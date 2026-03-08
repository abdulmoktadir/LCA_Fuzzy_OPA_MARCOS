import streamlit as st
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from pulp import LpMaximize, LpProblem, LpVariable, lpSum, value

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="Integrated Fuzzy MCDM Models",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =========================================================
# CUSTOM CSS
# =========================================================
st.markdown(
    """
<style>
:root {
    --primary: #1f77b4;
    --secondary: #2ca02c;
    --accent: #ff6b6b;
    --background: #f8f9fa;
    --card-bg: #ffffff;
    --text: #262730;
    --border-radius: 12px;
    --shadow: 0 6px 16px rgba(0, 0, 0, 0.08);
}

.main-header {
    font-size: 2.6rem;
    color: var(--primary);
    text-align: center;
    padding: 1.2rem 0;
    font-weight: 700;
    margin-bottom: 0.5rem;
    background: linear-gradient(135deg, var(--primary), var(--secondary));
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.section-header {
    font-size: 1.5rem;
    color: var(--primary);
    border-left: 5px solid var(--secondary);
    padding-left: 1rem;
    margin: 1.5rem 0 1rem 0;
    font-weight: 600;
}

.panel {
    background-color: var(--card-bg);
    border-radius: var(--border-radius);
    padding: 1.5rem;
    box-shadow: var(--shadow);
    margin-bottom: 1.5rem;
    border: 1px solid #e0e0e0;
}

.result-table {
    background-color: var(--card-bg);
    border-radius: var(--border-radius);
    padding: 1.2rem;
    box-shadow: var(--shadow);
    margin: 1rem 0;
}

.metric-card {
    background: linear-gradient(135deg, var(--card-bg), #f7f9fc);
    border-radius: var(--border-radius);
    padding: 1.2rem;
    box-shadow: var(--shadow);
    text-align: center;
    border: 1px solid #e8e8e8;
}

.metric-value {
    font-size: 1.7rem;
    font-weight: 700;
    color: var(--primary);
    margin: 0.5rem 0;
}

.metric-label {
    font-size: 0.9rem;
    color: #666;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

.optimization-formulation {
    background-color: #f8f9fa;
    border-left: 4px solid var(--primary);
    padding: 1.2rem;
    border-radius: 8px;
    margin: 1rem 0;
    font-family: 'Courier New', monospace;
    font-size: 0.92rem;
    line-height: 1.6;
}

.instruction-box {
    background-color: #f0f7ff;
    border-left: 4px solid var(--primary);
    padding: 1rem 1.2rem;
    border-radius: 4px;
    margin: 1rem 0;
}

.footer {
    text-align: center;
    margin-top: 2rem;
    padding: 1rem;
    color: #666;
    font-size: 0.9rem;
    border-top: 1px solid #eaeaea;
}

.stButton > button {
    background: linear-gradient(135deg, var(--primary), var(--secondary));
    color: white;
    border: none;
    padding: 0.75rem 1.5rem;
    border-radius: 40px;
    font-weight: 600;
    width: 100%;
}
</style>
    """,
    unsafe_allow_html=True,
)

# =========================================================
# TRIANGULAR FUZZY SCALE
# =========================================================
# Format: (l, m, u)
LINGUISTIC_SCALE = {
    "AL": (1.00, 1.50, 2.50),   # Absolutely Low
    "VL": (1.50, 2.50, 3.50),   # Very Low
    "L":  (2.50, 3.50, 4.50),   # Low
    "ML": (3.50, 4.50, 5.50),   # Medium Low
    "E":  (4.50, 5.50, 6.50),   # Equal
    "MH": (5.50, 6.50, 7.50),   # Medium High
    "H":  (6.50, 7.50, 8.50),   # High
    "VH": (7.50, 8.50, 9.50),   # Very High
    "AH": (8.50, 9.00, 10.00),  # Absolutely High
}

LINGUISTIC_LABELS = {
    "AL": "Absolutely Low",
    "VL": "Very Low",
    "L": "Low",
    "ML": "Medium Low",
    "E": "Equal",
    "MH": "Medium High",
    "H": "High",
    "VH": "Very High",
    "AH": "Absolutely High",
}

LINGUISTIC_OPTIONS = list(LINGUISTIC_SCALE.keys())


# =========================================================
# FUZZY FUNCTIONS
# =========================================================
def trig_geom_component(values, weights):
    """
    Trigonometric aggregation component matching the user's Excel formula:
    aggregated = sum(values) * (2/pi) * acos( product( cos(pi * v/sum(values) / 2) ** w ) )
    """
    s = sum(values)
    if s == 0:
        return 0.0

    prod = 1.0
    for v, w in zip(values, weights):
        ratio = v / s
        term = np.cos(np.pi * ratio / 2) ** w
        prod *= term

    prod = np.clip(prod, -1.0, 1.0)
    return s * (2 / np.pi) * np.arccos(prod)


def aggregate_tfn(tfn_list, weights):
    ls = [t[0] for t in tfn_list]
    ms = [t[1] for t in tfn_list]
    us = [t[2] for t in tfn_list]
    l = trig_geom_component(ls, weights)
    m = trig_geom_component(ms, weights)
    u = trig_geom_component(us, weights)
    return (l, m, u)


def defuzz_tfn(tfn):
    """Graded mean integration representation for TFN."""
    return (tfn[0] + 4 * tfn[1] + tfn[2]) / 6.0


def crisp_to_tfn(x, alpha=0.05):
    x = float(x)
    return (x, x + alpha, x + 2 * alpha)


def format_tfn(tfn):
    return f"({tfn[0]:.4f}, {tfn[1]:.4f}, {tfn[2]:.4f})"


# =========================================================
# TRIANGULAR FUZZY OPA
# =========================================================
def solve_fuzzy_opa(coeff_list, n):
    prob = LpProblem("Triangular_Fuzzy_OPA", LpMaximize)

    w_l = [LpVariable(f"w_l_{i}", lowBound=0) for i in range(n)]
    w_m = [LpVariable(f"w_m_{i}", lowBound=0) for i in range(n)]
    w_u = [LpVariable(f"w_u_{i}", lowBound=0) for i in range(n)]

    psi_l = LpVariable("psi_l", lowBound=0)
    psi_m = LpVariable("psi_m", lowBound=0)
    psi_u = LpVariable("psi_u", lowBound=0)

    prob += (psi_l + 2 * psi_m + psi_u) / 4

    for i in range(n):
        prob += w_l[i] <= w_m[i]
        prob += w_m[i] <= w_u[i]

    # TFN normalization
    prob += lpSum(w_l) == 0.8
    prob += lpSum(w_m) == 1.0
    prob += lpSum(w_u) == 1.2

    for a in range(n - 1):
        prob += coeff_list[a][0] * (w_l[a] - w_u[a + 1]) >= psi_l
        prob += coeff_list[a][1] * (w_m[a] - w_m[a + 1]) >= psi_m
        prob += coeff_list[a][2] * (w_u[a] - w_l[a + 1]) >= psi_u

    prob += coeff_list[n - 1][0] * w_l[n - 1] >= psi_l
    prob += coeff_list[n - 1][1] * w_m[n - 1] >= psi_m
    prob += coeff_list[n - 1][2] * w_u[n - 1] >= psi_u

    status = prob.solve()
    if status != 1:
        st.error("Optimization failed. The model may be infeasible. Please check your inputs.")
        return None, None

    weights = []
    for i in range(n):
        weights.append((
            max(0, value(w_l[i])),
            max(0, value(w_m[i])),
            max(0, value(w_u[i]))
        ))

    psi = (
        max(0, value(psi_l)),
        max(0, value(psi_m)),
        max(0, value(psi_u))
    )
    return weights, psi


def display_linguistic_scale():
    scale_df = pd.DataFrame([
        {"Code": k, "Linguistic Attribute": LINGUISTIC_LABELS[k], "l": v[0], "m": v[1], "u": v[2]}
        for k, v in LINGUISTIC_SCALE.items()
    ])
    st.dataframe(scale_df, use_container_width=True, hide_index=True)


def create_opa_word_document(criteria, theta, defuzz_values, coeff, ranked_criteria, weights, psi, num_experts, expert_weights):
    doc = Document()
    title = doc.add_heading('Triangular Fuzzy OPA Results - Multiple Experts', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'This document contains the results of the Triangular Fuzzy OPA analysis for {num_experts} experts.')
    doc.add_paragraph('')

    doc.add_heading('Expert Weights', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Expert'
    hdr[1].text = 'Weight'
    for e in range(num_experts):
        row = table.add_row().cells
        row[0].text = f'Expert {e+1}'
        row[1].text = f'{expert_weights[e]:.6f}'

    doc.add_paragraph('')
    doc.add_heading('Aggregated Triangular Fuzzy Importance (Theta)', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Criterion'
    hdr[1].text = 'l'
    hdr[2].text = 'm'
    hdr[3].text = 'u'
    hdr[4].text = 'Defuzzified'
    for i, crit in enumerate(criteria):
        row = table.add_row().cells
        row[0].text = crit
        row[1].text = f'{theta[i][0]:.4f}'
        row[2].text = f'{theta[i][1]:.4f}'
        row[3].text = f'{theta[i][2]:.4f}'
        row[4].text = f'{defuzz_values[i]:.4f}'

    doc.add_paragraph('')
    doc.add_heading('Coefficients for Fuzzy OPA', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Criterion'
    hdr[1].text = 'Coeff l'
    hdr[2].text = 'Coeff m'
    hdr[3].text = 'Coeff u'
    for i, crit in enumerate(criteria):
        row = table.add_row().cells
        row[0].text = crit
        row[1].text = f'{coeff[i][0]:.4f}'
        row[2].text = f'{coeff[i][1]:.4f}'
        row[3].text = f'{coeff[i][2]:.4f}'

    doc.add_paragraph('')
    doc.add_heading('Ranked Criteria', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Rank'
    hdr[1].text = 'Criterion'
    hdr[2].text = 'Weight l'
    hdr[3].text = 'Weight m'
    hdr[4].text = 'Weight u'
    hdr[5].text = 'Defuzzified'
    for rank, (crit_idx, w) in enumerate(ranked_criteria):
        row = table.add_row().cells
        row[0].text = str(rank + 1)
        row[1].text = criteria[crit_idx]
        row[2].text = f'{w[0]:.4f}'
        row[3].text = f'{w[1]:.4f}'
        row[4].text = f'{w[2]:.4f}'
        row[5].text = f'{defuzz_tfn(w):.4f}'

    doc.add_paragraph('')
    doc.add_heading('Psi Value', level=1)
    doc.add_paragraph(f'Psi: ({psi[0]:.4f}, {psi[1]:.4f}, {psi[2]:.4f})')
    doc.add_paragraph(f'Defuzzified Psi: {defuzz_tfn(psi):.4f}')

    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


def display_optimization_formulation(coeff_list, criteria_names):
    st.markdown('<div class="optimization-formulation">', unsafe_allow_html=True)
    st.markdown(
        """
        <strong>Triangular Fuzzy OPA Formulation</strong><br><br>
        Objective:<br>
        Maximize: (Ψ_l + 2*Ψ_m + Ψ_u) / 4<br><br>
        Constraints:<br>
        1. Triangular ordering: w_lᵢ ≤ w_mᵢ ≤ w_uᵢ<br>
        2. Normalization: Σw_lᵢ = 0.8, Σw_mᵢ = 1.0, Σw_uᵢ = 1.2<br>
        3. Adjacent ranked criteria:<br>
        &nbsp;&nbsp;θ_lᵃ(w_lᵃ - w_uᵃ⁺¹) ≥ Ψ_l<br>
        &nbsp;&nbsp;θ_mᵃ(w_mᵃ - w_mᵃ⁺¹) ≥ Ψ_m<br>
        &nbsp;&nbsp;θ_uᵃ(w_uᵃ - w_lᵃ⁺¹) ≥ Ψ_u<br>
        4. Last criterion:<br>
        &nbsp;&nbsp;θ_lⁿw_lⁿ ≥ Ψ_l, θ_mⁿw_mⁿ ≥ Ψ_m, θ_uⁿw_uⁿ ≥ Ψ_u
        """,
        unsafe_allow_html=True,
    )
    st.markdown('</div>', unsafe_allow_html=True)

    coeff_df = pd.DataFrame({
        'Criterion': criteria_names,
        'θ_l': [f'{c[0]:.4f}' for c in coeff_list],
        'θ_m': [f'{c[1]:.4f}' for c in coeff_list],
        'θ_u': [f'{c[2]:.4f}' for c in coeff_list],
    })
    st.dataframe(coeff_df, use_container_width=True, hide_index=True)


def opa_model():
    st.markdown('<h1 class="main-header">Triangular Fuzzy OPA Analysis with Multiple Experts</h1>', unsafe_allow_html=True)
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 2rem; color: #555;">
        This application implements a <b>Triangular Fuzzy OPA</b> model using your updated linguistic scale.
        </div>
        """,
        unsafe_allow_html=True,
    )

    if 'opa_criteria' not in st.session_state:
        st.session_state.opa_criteria = [f'Criterion {i+1}' for i in range(5)]
    if 'opa_num_criteria' not in st.session_state:
        st.session_state.opa_num_criteria = 5
    if 'opa_num_experts' not in st.session_state:
        st.session_state.opa_num_experts = 2
    if 'opa_results_calculated' not in st.session_state:
        st.session_state.opa_results_calculated = False
    if 'opa_expert_data' not in st.session_state:
        st.session_state.opa_expert_data = {}
    if 'opa_expert_weights' not in st.session_state:
        st.session_state.opa_expert_weights = [0.5, 0.5]

    left_col, right_col = st.columns([1, 1])

    with left_col:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.markdown('<h2 class="section-header">Step 1: Define Criteria and Experts</h2>', unsafe_allow_html=True)

        num_experts = st.number_input('Number of experts', min_value=1, max_value=15, value=st.session_state.opa_num_experts, step=1)
        num_criteria = st.number_input('Number of criteria', min_value=3, max_value=25, value=st.session_state.opa_num_criteria, step=1)

        if num_criteria != st.session_state.opa_num_criteria or num_experts != st.session_state.opa_num_experts:
            st.session_state.opa_num_criteria = num_criteria
            st.session_state.opa_num_experts = num_experts
            if len(st.session_state.opa_criteria) < num_criteria:
                for i in range(len(st.session_state.opa_criteria), num_criteria):
                    st.session_state.opa_criteria.append(f'Criterion {i+1}')
            else:
                st.session_state.opa_criteria = st.session_state.opa_criteria[:num_criteria]
            st.session_state.opa_expert_data = {}
            st.session_state.opa_expert_weights = [1.0 / num_experts] * num_experts
            st.session_state.opa_results_calculated = False

        criteria = []
        for i in range(num_criteria):
            c = st.text_input(f'Criterion {i+1}', value=st.session_state.opa_criteria[i], key=f'opa_criterion_{i}')
            criteria.append(c)
        st.session_state.opa_criteria = criteria

        st.markdown('<h3>Expert Weights</h3>', unsafe_allow_html=True)
        expert_weights = []
        for e in range(num_experts):
            w = st.number_input(
                f'Weight for Expert {e+1}',
                min_value=0.0,
                max_value=1.0,
                value=float(st.session_state.opa_expert_weights[e]) if e < len(st.session_state.opa_expert_weights) else 1.0 / num_experts,
                step=0.000001,
                format='%.6f',
                key=f'opa_expert_w_{e}'
            )
            expert_weights.append(w)

        sum_w = sum(expert_weights)
        st.write(f'**Sum of expert weights: {sum_w:.6f}**')
        if abs(sum_w - 1.0) > 1e-6:
            st.error('Expert weights must sum to 1.0.')
            st.session_state.opa_weights_valid = False
        else:
            st.session_state.opa_expert_weights = expert_weights
            st.session_state.opa_weights_valid = True
            st.success('Expert weights are valid.')

        st.markdown('</div>', unsafe_allow_html=True)

    with right_col:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.markdown('<h2 class="section-header">Expert Assessments</h2>', unsafe_allow_html=True)
        st.markdown('<div class="instruction-box"><b>Updated linguistic scale</b></div>', unsafe_allow_html=True)
        display_linguistic_scale()

        tabs = st.tabs([f'Expert {e+1}' for e in range(num_experts)])
        for e, tab in enumerate(tabs):
            with tab:
                if f'expert_{e}' not in st.session_state.opa_expert_data:
                    st.session_state.opa_expert_data[f'expert_{e}'] = ['E'] * num_criteria

                df = pd.DataFrame({
                    'Criterion': criteria,
                    'Rating': st.session_state.opa_expert_data[f'expert_{e}']
                })
                edited_df = st.data_editor(
                    df,
                    column_config={
                        'Rating': st.column_config.SelectboxColumn('Rating', options=LINGUISTIC_OPTIONS, required=True)
                    },
                    use_container_width=True,
                    key=f'opa_data_editor_expert_{e}'
                )
                st.session_state.opa_expert_data[f'expert_{e}'] = edited_df['Rating'].tolist()

        all_ratings_available = all(f'expert_{e}' in st.session_state.opa_expert_data for e in range(num_experts))
        weights_valid = getattr(st.session_state, 'opa_weights_valid', False)
        disabled = not (all_ratings_available and weights_valid)

        if st.button('Calculate Weights', disabled=disabled, use_container_width=True):
            theta = []
            defuzz_values = []
            for j in range(num_criteria):
                tfn_list = [LINGUISTIC_SCALE[st.session_state.opa_expert_data[f'expert_{e}'][j]] for e in range(num_experts)]
                aggregated = aggregate_tfn(tfn_list, st.session_state.opa_expert_weights)
                theta.append(aggregated)
                defuzz_values.append(defuzz_tfn(aggregated))

            min_l = min(t[0] for t in theta if t[0] > 0) if any(t[0] > 0 for t in theta) else 1.0

            coeff = []
            for t in theta:
                if min(t) <= 0:
                    coeff.append((0, 0, 0))
                else:
                    coeff.append((min_l / t[2], min_l / t[1], min_l / t[0]))

            sorted_indices = np.argsort(defuzz_values)[::-1]
            coeff_sorted = [coeff[idx] for idx in sorted_indices]
            weights_sorted, psi = solve_fuzzy_opa(coeff_sorted, num_criteria)

            if weights_sorted is not None:
                weights = [None] * num_criteria
                for rank, idx in enumerate(sorted_indices):
                    weights[idx] = weights_sorted[rank]
                ranked_criteria = [(sorted_indices[k], weights_sorted[k]) for k in range(num_criteria)]

                st.session_state.opa_theta = theta
                st.session_state.opa_defuzz_values = defuzz_values
                st.session_state.opa_coeff = coeff
                st.session_state.opa_ranked_criteria = ranked_criteria
                st.session_state.opa_weights = weights
                st.session_state.opa_psi = psi
                st.session_state.opa_results_calculated = True
                st.success('Calculation completed successfully.')

        st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.opa_results_calculated:
        st.markdown('<h2 class="section-header">Results</h2>', unsafe_allow_html=True)

        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader('Normalized Expert Weights')
        df_expert_w = pd.DataFrame({
            'Expert': [f'Expert {e+1}' for e in range(st.session_state.opa_num_experts)],
            'Weight': st.session_state.opa_expert_weights
        })
        st.dataframe(df_expert_w, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader('Aggregated Triangular Fuzzy Importance (Theta)')
        df_theta = pd.DataFrame({
            'Criterion': st.session_state.opa_criteria,
            'l': [t[0] for t in st.session_state.opa_theta],
            'm': [t[1] for t in st.session_state.opa_theta],
            'u': [t[2] for t in st.session_state.opa_theta],
            'Defuzzified': st.session_state.opa_defuzz_values,
        })
        st.dataframe(df_theta, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader('Coefficients for Fuzzy OPA')
        df_coeff = pd.DataFrame({
            'Criterion': st.session_state.opa_criteria,
            'Coeff l': [c[0] for c in st.session_state.opa_coeff],
            'Coeff m': [c[1] for c in st.session_state.opa_coeff],
            'Coeff u': [c[2] for c in st.session_state.opa_coeff],
        })
        st.dataframe(df_coeff, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        sorted_indices = np.argsort(st.session_state.opa_defuzz_values)[::-1]
        coeff_sorted = [st.session_state.opa_coeff[idx] for idx in sorted_indices]
        criteria_sorted = [st.session_state.opa_criteria[idx] for idx in sorted_indices]
        display_optimization_formulation(coeff_sorted, criteria_sorted)

        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader('Triangular Fuzzy Weights')
        df_weights = pd.DataFrame({
            'Criterion': st.session_state.opa_criteria,
            'l': [w[0] for w in st.session_state.opa_weights],
            'm': [w[1] for w in st.session_state.opa_weights],
            'u': [w[2] for w in st.session_state.opa_weights],
            'Defuzzified': [defuzz_tfn(w) for w in st.session_state.opa_weights],
        })
        st.dataframe(df_weights, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="result-table">', unsafe_allow_html=True)
        st.subheader('Ranked Criteria and Weights')
        df_ranked = pd.DataFrame({
            'Rank': list(range(1, st.session_state.opa_num_criteria + 1)),
            'Criterion': [st.session_state.opa_criteria[idx] for idx, _ in st.session_state.opa_ranked_criteria],
            'Weight l': [w[0] for _, w in st.session_state.opa_ranked_criteria],
            'Weight m': [w[1] for _, w in st.session_state.opa_ranked_criteria],
            'Weight u': [w[2] for _, w in st.session_state.opa_ranked_criteria],
            'Defuzzified': [defuzz_tfn(w) for _, w in st.session_state.opa_ranked_criteria],
        })
        st.dataframe(df_ranked, use_container_width=True, hide_index=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.markdown('<div class="metric-label">Psi (l, m, u)</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="metric-value">{st.session_state.opa_psi[0]:.4f}, {st.session_state.opa_psi[1]:.4f}, {st.session_state.opa_psi[2]:.4f}</div>', unsafe_allow_html=True)
        st.markdown('<div class="metric-label">Defuzzified Psi</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="metric-value">{defuzz_tfn(st.session_state.opa_psi):.4f}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        doc_bytes = create_opa_word_document(
            st.session_state.opa_criteria,
            st.session_state.opa_theta,
            st.session_state.opa_defuzz_values,
            st.session_state.opa_coeff,
            st.session_state.opa_ranked_criteria,
            st.session_state.opa_weights,
            st.session_state.opa_psi,
            st.session_state.opa_num_experts,
            st.session_state.opa_expert_weights,
        )
        st.download_button(
            label='Export Results to Word',
            data=doc_bytes,
            file_name='Triangular_Fuzzy_OPA_Results.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            use_container_width=True,
        )

    with st.expander('Learn more about the Triangular Fuzzy OPA Method'):
        st.markdown(
            """
            1. Experts provide linguistic assessments.
            2. Linguistic terms are converted into triangular fuzzy numbers (TFNs).
            3. Expert opinions are aggregated using the trigonometric fuzzy weighted geometric operator.
            4. Aggregated TFNs are defuzzified and ranked.
            5. A triangular fuzzy linear programming model computes final weights.
            """
        )
        display_linguistic_scale()


# =========================================================
# TRIANGULAR FUZZY MARCOS
# =========================================================
def tfn_add(a, b):
    return (a[0] + b[0], a[1] + b[1], a[2] + b[2])


def tfn_div(a, b):
    eps = 1e-9
    return (
        a[0] / max(b[2], eps),
        a[1] / max(b[1], eps),
        a[2] / max(b[0], eps),
    )


def tfn_scalar_div(a, scalar):
    scalar = max(float(scalar), 1e-9)
    return (a[0] / scalar, a[1] / scalar, a[2] / scalar)


def defuzz_marcos_tfn(a):
    return (a[0] + 4 * a[1] + a[2]) / 6.0


def generalized_mean_tfn(a):
    l, m, u = float(a[0]), float(a[1]), float(a[2])
    # Match Excel exactly: IF(u>0, formula, 0)
    if u <= 0:
        return 0.0
    denom = 3.0 * (u - l)
    if abs(denom) < 1e-12:
        return 0.0
    return (u * u + m * u - l * m - l * l) / denom


def crisp_to_tfn_10pct(x):
    x = float(x)
    if x == 0:
        return (0.0, 0.0, 0.0)
    if x > 0:
        return (0.9 * x, x, 1.1 * x)
    return (1.1 * x, x, 0.9 * x)


def create_marcos_word_document(all_data):
    doc = Document()
    title = doc.add_heading('Triangular Fuzzy MARCOS Results', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Problem Setup', level=1)
    setup = doc.add_table(rows=5, cols=2)
    setup.style = 'Table Grid'
    rows = [
        ('Number of Alternatives', all_data['n_alternatives']),
        ('Number of Criteria', all_data['n_criteria']),
        ('Number of Experts', all_data['n_experts']),
        ('Hard Criteria Uncertainty', '10%'),
        ('Defuzzification', '(l + 4m + u)/6 at final stage'),
    ]
    for i, (k, v) in enumerate(rows):
        setup.rows[i].cells[0].text = str(k)
        setup.rows[i].cells[1].text = str(v)

    doc.add_paragraph('')
    doc.add_heading('Final Ranking Results', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Rank'
    hdr[1].text = 'Alternative'
    hdr[2].text = 'Crisp Ki-'
    hdr[3].text = 'Crisp Ki+'
    hdr[4].text = 'Utility f(Ki)'
    hdr[5].text = 'Order'

    for _, row_data in all_data['final_results'].iterrows():
        row = table.add_row().cells
        row[0].text = str(int(row_data['Rank']))
        row[1].text = str(row_data['Alternative'])
        row[2].text = f"{row_data['Crisp Ki-']:.6f}"
        row[3].text = f"{row_data['Crisp Ki+']:.6f}"
        row[4].text = f"{row_data['f(Ki)']:.6f}"
        row[5].text = str(int(row_data['Rank']))

    doc.add_paragraph('')
    doc.add_paragraph(f"Best Alternative: {all_data['best_alternative']} with utility score {all_data['best_score']:.6f}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def marcos_model():
    st.markdown('<h1 class="main-header">Triangular Fuzzy MARCOS Method</h1>', unsafe_allow_html=True)
    st.write('Soft criteria are aggregated using the trigonometric operator, hard criteria are converted into TFNs using 10% uncertainty, and defuzzification is performed at the final MARCOS stage.')

    with st.expander('Learn more about the Triangular Fuzzy MARCOS Method', expanded=False):
        st.markdown(
            """
            **Implemented logic in this module**
            - Soft criteria: expert linguistic ratings → TFNs → trigonometric aggregation.
            - Hard criteria: crisp values → TFNs with **10% uncertainty**.
            - MARCOS computations are carried out in fuzzy form.
            - **Defuzzification is postponed to the final stage**.
            """
        )
        display_linguistic_scale()

    if 'marcos_step' not in st.session_state:
        st.session_state.marcos_step = 1
    if 'marcos_data' not in st.session_state:
        st.session_state.marcos_data = {}

    steps = [
        'Problem Setup', 'Criteria Setup', 'Expert Weights', 'Data Collection',
        'Build Fuzzy Decision Matrix', 'Criteria Information', 'Results'
    ]
    current_step = st.session_state.marcos_step - 1
    st.progress(current_step / (len(steps) - 1))
    st.write(f'**Current Step: {steps[current_step]}**')

    if st.session_state.marcos_step == 1:
        marcos_step1_input()
    elif st.session_state.marcos_step == 2:
        marcos_step2_criteria_setup()
    elif st.session_state.marcos_step == 3:
        marcos_step3_expert_weights()
    elif st.session_state.marcos_step == 4:
        marcos_step4_data_collection()
    elif st.session_state.marcos_step == 5:
        marcos_step5_decision_matrix()
    elif st.session_state.marcos_step == 6:
        marcos_step6_criteria_info()
    elif st.session_state.marcos_step == 7:
        marcos_step7_calculations()


def marcos_step1_input():
    st.header('Step 1: Problem Setup')
    n_alternatives = st.number_input('Number of alternatives', min_value=2, max_value=50, value=4, key='marcos_n_alts')
    n_criteria = st.number_input('Number of criteria', min_value=2, max_value=25, value=6, key='marcos_n_criteria')
    n_experts = st.number_input('Number of experts', min_value=1, max_value=15, value=3, key='marcos_n_experts')

    if st.button('Next: Criteria Setup', key='marcos_to_step2'):
        st.session_state.marcos_data['n_alternatives'] = int(n_alternatives)
        st.session_state.marcos_data['n_criteria'] = int(n_criteria)
        st.session_state.marcos_data['n_experts'] = int(n_experts)
        st.session_state.marcos_data['alternatives'] = [f'A{i+1}' for i in range(int(n_alternatives))]
        st.session_state.marcos_data['criteria'] = [f'C{i+1}' for i in range(int(n_criteria))]
        st.session_state.marcos_data['criteria_types'] = ['Soft'] * int(n_criteria)
        st.session_state.marcos_data['expert_weights'] = [1.0 / int(n_experts)] * int(n_experts)
        st.session_state.marcos_step = 2
        st.rerun()


def marcos_step2_criteria_setup():
    st.header('Step 2: Criteria Setup')
    n_criteria = st.session_state.marcos_data['n_criteria']
    criteria = st.session_state.marcos_data['criteria']

    df = pd.DataFrame({
        'Criterion': criteria,
        'Type': st.session_state.marcos_data['criteria_types'],
        'Description': [''] * n_criteria,
    })
    edited_df = st.data_editor(
        df,
        use_container_width=True,
        column_config={
            'Type': st.column_config.SelectboxColumn('Type', options=['Soft', 'Hard'])
        },
        key='marcos_criteria_setup'
    )

    with st.expander('Linguistic Scale for Soft Criteria'):
        display_linguistic_scale()

    if st.button('Next: Expert Weights', key='marcos_to_step3'):
        st.session_state.marcos_data['criteria_setup'] = edited_df
        st.session_state.marcos_data['criteria_types'] = edited_df['Type'].tolist()
        st.session_state.marcos_step = 3
        st.rerun()


def marcos_step3_expert_weights():
    st.header('Step 3: Expert Weights')
    n_experts = st.session_state.marcos_data['n_experts']
    expert_weights = []

    for e in range(n_experts):
        weight = st.number_input(
            f'Weight for Expert {e+1}',
            min_value=0.0,
            max_value=1.0,
            value=float(st.session_state.marcos_data['expert_weights'][e]),
            step=0.0001,
            format='%.4f',
            key=f'marcos_expert_weight_{e}'
        )
        expert_weights.append(weight)

    weight_sum = sum(expert_weights)
    valid_weights = abs(weight_sum - 1.0) <= 1e-6
    if valid_weights:
        st.success('Expert weights are valid.')
    else:
        st.warning(f'Expert weights sum to {weight_sum:.4f}, but should sum to 1.0')

    if st.button('Next: Data Collection', key='marcos_to_step4') and valid_weights:
        st.session_state.marcos_data['expert_weights'] = expert_weights
        st.session_state.marcos_step = 4
        st.rerun()


def marcos_step4_data_collection():
    st.header('Step 4: Data Collection')
    n_alternatives = st.session_state.marcos_data['n_alternatives']
    n_criteria = st.session_state.marcos_data['n_criteria']
    n_experts = st.session_state.marcos_data['n_experts']
    alternatives = st.session_state.marcos_data['alternatives']
    criteria = st.session_state.marcos_data['criteria']
    criteria_types = st.session_state.marcos_data['criteria_types']

    if 'expert_data' not in st.session_state.marcos_data:
        st.session_state.marcos_data['expert_data'] = {}
        for e in range(n_experts):
            st.session_state.marcos_data['expert_data'][f'expert_{e}'] = pd.DataFrame(
                [['E'] * n_criteria for _ in range(n_alternatives)],
                columns=criteria,
                index=alternatives,
            )

    if 'hard_data' not in st.session_state.marcos_data:
        st.session_state.marcos_data['hard_data'] = pd.DataFrame(
            [[0.0] * n_criteria for _ in range(n_alternatives)],
            columns=criteria,
            index=alternatives,
        )

    soft_criteria = [criteria[i] for i, t in enumerate(criteria_types) if str(t).lower() == 'soft']
    hard_criteria = [criteria[i] for i, t in enumerate(criteria_types) if str(t).lower() == 'hard']

    st.write('### Soft Criteria Assessment')
    if soft_criteria:
        expert_tabs = st.tabs([f'Expert {e+1}' for e in range(n_experts)])
        for e, tab in enumerate(expert_tabs):
            with tab:
                current_data = st.session_state.marcos_data['expert_data'][f'expert_{e}']
                soft_data = current_data[soft_criteria].copy()
                edited_soft = st.data_editor(
                    soft_data,
                    column_config={
                        col: st.column_config.SelectboxColumn(col, options=LINGUISTIC_OPTIONS, required=True)
                        for col in soft_criteria
                    },
                    use_container_width=True,
                    key=f'marcos_expert_{e}_soft'
                )
                st.session_state.marcos_data['expert_data'][f'expert_{e}'].update(edited_soft)

    st.write('### Hard Criteria Assessment')
    st.caption('Hard criteria will be converted to TFNs using ±10% uncertainty.')
    if hard_criteria:
        current_hard = st.session_state.marcos_data['hard_data'][hard_criteria].copy()
        edited_hard = st.data_editor(current_hard, use_container_width=True, key='marcos_hard_data')
        st.session_state.marcos_data['hard_data'].update(edited_hard)

    if st.button('Next: Build Fuzzy Decision Matrix', key='marcos_to_step5'):
        st.session_state.marcos_step = 5
        st.rerun()


def marcos_step5_decision_matrix():
    st.header('Step 5: Build Fuzzy Decision Matrix')
    n_alternatives = st.session_state.marcos_data['n_alternatives']
    n_experts = st.session_state.marcos_data['n_experts']
    alternatives = st.session_state.marcos_data['alternatives']
    criteria = st.session_state.marcos_data['criteria']
    criteria_types = st.session_state.marcos_data['criteria_types']
    expert_weights = st.session_state.marcos_data['expert_weights']

    fuzzy_matrix = [[None for _ in criteria] for _ in alternatives]
    detail_rows = []

    for i, alt in enumerate(alternatives):
        for j, criterion in enumerate(criteria):
            if str(criteria_types[j]).lower() == 'soft':
                tfn_list = []
                labels = []
                for e in range(n_experts):
                    val = st.session_state.marcos_data['expert_data'][f'expert_{e}'].loc[alt, criterion]
                    labels.append(val)
                    tfn_list.append(LINGUISTIC_SCALE[val])
                agg = aggregate_tfn(tfn_list, expert_weights)
                fuzzy_matrix[i][j] = agg
                detail_rows.append({
                    'Alternative': alt,
                    'Criterion': criterion,
                    'Type': 'Soft',
                    'Inputs': ', '.join(labels),
                    'Aggregated TFN': format_tfn(agg),
                })
            else:
                crisp = st.session_state.marcos_data['hard_data'].loc[alt, criterion]
                tfn = crisp_to_tfn_10pct(crisp)
                fuzzy_matrix[i][j] = tfn
                detail_rows.append({
                    'Alternative': alt,
                    'Criterion': criterion,
                    'Type': 'Hard',
                    'Inputs': f'{crisp}',
                    'Aggregated TFN': format_tfn(tfn),
                })

    st.subheader('Fuzzy Decision Matrix Details')
    st.dataframe(pd.DataFrame(detail_rows), use_container_width=True, hide_index=True)

    display_df = pd.DataFrame(index=alternatives, columns=criteria)
    for i, alt in enumerate(alternatives):
        for j, criterion in enumerate(criteria):
            display_df.loc[alt, criterion] = format_tfn(fuzzy_matrix[i][j])
    st.subheader('Final Fuzzy Decision Matrix')
    st.dataframe(display_df, use_container_width=True)

    st.session_state.marcos_data['fuzzy_matrix'] = fuzzy_matrix
    st.session_state.marcos_step = 6
    if st.button('Next: Criteria Information', key='marcos_to_step6'):
        st.rerun()


def marcos_step6_criteria_info():
    st.header('Step 6: Criteria Information')
    n_criteria = st.session_state.marcos_data['n_criteria']
    criteria = st.session_state.marcos_data['criteria']

    default_types = ['Benefit'] * n_criteria
    default_l = [1.0 / n_criteria] * n_criteria
    default_m = [1.0 / n_criteria] * n_criteria
    default_u = [1.0 / n_criteria] * n_criteria

    # If OPA results are available and dimensions match, use them as default fuzzy weights
    opa_ready = (
        'opa_weights' in st.session_state and
        isinstance(st.session_state.opa_weights, list) and
        len(st.session_state.opa_weights) == n_criteria
    )
    if opa_ready:
        default_l = [float(w[0]) for w in st.session_state.opa_weights]
        default_m = [float(w[1]) for w in st.session_state.opa_weights]
        default_u = [float(w[2]) for w in st.session_state.opa_weights]
        st.info('OPA fuzzy weights detected and used as default TFN weights for MARCOS.')

    df = pd.DataFrame({
        'Criterion': criteria,
        'Type': default_types,
        'Weight l': default_l,
        'Weight m': default_m,
        'Weight u': default_u,
    })

    edited_df = st.data_editor(
        df,
        use_container_width=True,
        column_config={
            'Type': st.column_config.SelectboxColumn('Type', options=['Benefit', 'Cost'])
        },
        key='marcos_criteria_info'
    )

    # Validation
    ordering_ok = True
    for _, row in edited_df.iterrows():
        if not (float(row['Weight l']) <= float(row['Weight m']) <= float(row['Weight u'])):
            ordering_ok = False
            break

    sum_l = edited_df['Weight l'].astype(float).sum()
    sum_m = edited_df['Weight m'].astype(float).sum()
    sum_u = edited_df['Weight u'].astype(float).sum()

    st.write(f"**Σ Weight l = {sum_l:.6f}, Σ Weight m = {sum_m:.6f}, Σ Weight u = {sum_u:.6f}**")

    if not ordering_ok:
        st.warning('Each fuzzy weight must satisfy: Weight l ≤ Weight m ≤ Weight u')
    else:
        st.success('TFN ordering of fuzzy weights is valid.')

    if st.button('Calculate MARCOS Results', key='marcos_to_step7') and ordering_ok:
        st.session_state.marcos_data['criteria_info'] = edited_df
        st.session_state.marcos_step = 7
        st.rerun()


def marcos_step7_calculations():
    st.header('Step 7: Fuzzy MARCOS Results')

    alternatives = st.session_state.marcos_data['alternatives']
    criteria = st.session_state.marcos_data['criteria']
    fuzzy_matrix = st.session_state.marcos_data['fuzzy_matrix']
    criteria_info = st.session_state.marcos_data['criteria_info']
    crit_types = criteria_info['Type'].tolist()
    fuzzy_weights = [
        (
            float(criteria_info.iloc[j]['Weight l']),
            float(criteria_info.iloc[j]['Weight m']),
            float(criteria_info.iloc[j]['Weight u'])
        )
        for j in range(len(criteria))
    ]

    n_alt = len(alternatives)
    n_crit = len(criteria)

    st.subheader('Fuzzy Criteria Weights')
    df_w = pd.DataFrame({
        'Criterion': criteria,
        'Type': crit_types,
        'Weight TFN': [format_tfn(w) for w in fuzzy_weights],
    })
    st.dataframe(df_w, use_container_width=True, hide_index=True)

    # =====================================================
    # Step 1: Generalized mean M(Vij) from integrated matrix
    # M(Vij)=(-a^2+c^2-ab+bc)/(3(-a+c))
    # =====================================================
    m_values = [[generalized_mean_tfn(fuzzy_matrix[i][j]) for j in range(n_crit)] for i in range(n_alt)]

    m_df = pd.DataFrame(m_values, index=alternatives, columns=criteria)
    st.subheader('Step 1: M(Vij) Matrix')
    st.dataframe(m_df, use_container_width=True)

    # =====================================================
    # Step 2: Determine A (AI) and A (ID) using M(Vij)
    # It uses scalar min/max of M(Vij) and repeats the value as (m,m,m).
    # Benefit: A (AI)=min M(Vij), A (ID)=max M(Vij)
    # Cost:    A (AI)=max M(Vij), A (ID)=min M(Vij)
    # =====================================================
    anti_ideal = []
    ideal = []
    ref_rows = []
    for j in range(n_crit):
        col_m = [m_values[i][j] for i in range(n_alt)]
        min_m = min(col_m)
        max_m = max(col_m)

        if crit_types[j] == 'Benefit':
            ai_scalar = min_m
            id_scalar = max_m
        else:
            ai_scalar = max_m
            id_scalar = min_m

        anti_ideal.append((ai_scalar, ai_scalar, ai_scalar))
        ideal.append((id_scalar, id_scalar, id_scalar))
        ref_rows.append({
            'Criterion': criteria[j],
            'Type': crit_types[j],
            'Min M(Vij)': min_m,
            'Max M(Vij)': max_m,
            'A (AI)': format_tfn((ai_scalar, ai_scalar, ai_scalar)),
            'A (ID)': format_tfn((id_scalar, id_scalar, id_scalar)),
        })

    st.subheader('Step 2: A (AI) and A (ID) from M(Vij)')
    st.dataframe(pd.DataFrame(ref_rows), use_container_width=True, hide_index=True)

    # =====================================================
    # Step 3: Extended integrated matrix
    # Row order follows Excel:
    # A (AI), A1..An, A (ID)
    # =====================================================
    labels = ['A (AI)'] + alternatives + ['A (ID)']
    extended = [anti_ideal] + fuzzy_matrix + [ideal]

    ext_df = pd.DataFrame(index=labels, columns=criteria)
    for i, lab in enumerate(labels):
        for j, c in enumerate(criteria):
            ext_df.loc[lab, c] = format_tfn(extended[i][j])
    st.subheader('Step 3: Extended Integrated Matrix')
    st.dataframe(ext_df, use_container_width=True)

    # =====================================================
    # Step 4: Normalized extended integrated matrix
    # Benefit: (x_l/ID_u, x_m/ID_u, x_u/ID_u)
    # Cost:    (ID_l/x_u, ID_l/x_m, ID_l/x_l)
    # =====================================================
    norm = [[None for _ in range(n_crit)] for _ in range(n_alt + 2)]
    for i in range(n_alt + 2):
        for j in range(n_crit):
            x = extended[i][j]
            if crit_types[j] == 'Benefit':
                denom = max(ideal[j][2], 1e-9)
                norm[i][j] = (
                    x[0] / denom,
                    x[1] / denom,
                    x[2] / denom,
                )
            else:
                a = ideal[j][0]
                norm[i][j] = (
                    a / max(x[2], 1e-9),
                    a / max(x[1], 1e-9),
                    a / max(x[0], 1e-9),
                )

    norm_df = pd.DataFrame(index=labels, columns=criteria)
    for i, lab in enumerate(labels):
        for j, c in enumerate(criteria):
            norm_df.loc[lab, c] = format_tfn(norm[i][j])
    st.subheader('Step 4: Normalized Extended Integrated Matrix')
    st.dataframe(norm_df, use_container_width=True)

    # =====================================================
    # Step 5: Weighted normalized matrix
    # =====================================================
    weighted = [[None for _ in range(n_crit)] for _ in range(n_alt + 2)]
    for i in range(n_alt + 2):
        for j in range(n_crit):
            weighted[i][j] = (
                norm[i][j][0] * fuzzy_weights[j][0],
                norm[i][j][1] * fuzzy_weights[j][1],
                norm[i][j][2] * fuzzy_weights[j][2],
            )

    weighted_df = pd.DataFrame(index=labels, columns=criteria)
    for i, lab in enumerate(labels):
        for j, c in enumerate(criteria):
            weighted_df.loc[lab, c] = format_tfn(weighted[i][j])
    st.subheader('Step 5: Weighted Normalized Matrix')
    st.dataframe(weighted_df, use_container_width=True)

    # =====================================================
    # Step 6: Si values (sum of TFNs across criteria)
    # =====================================================
    S = []
    for i in range(n_alt + 2):
        s = (0.0, 0.0, 0.0)
        for j in range(n_crit):
            s = tfn_add(s, weighted[i][j])
        S.append(s)

    s_df = pd.DataFrame({
        'Alternative': labels,
        'S_i (TFN)': [format_tfn(x) for x in S],
        'Crisp T_i': [defuzz_marcos_tfn(x) for x in S],
    })
    st.subheader('Step 6: Aggregated Fuzzy Utility (S_i)')
    st.dataframe(s_df, use_container_width=True, hide_index=True)

    S_ai = S[0]   # A (AI)
    S_id = S[-1]  # A (ID)

    # =====================================================
    # Step 7: Fuzzy Ki- and Ki+
    # Ki- = Si / S(AI)
    # Ki+ = Si / S(ID)
    # with triangular fuzzy division ordering:
    # (l/u, m/m, u/l)
    # =====================================================
    results = []
    for idx, alt in enumerate(alternatives, start=1):
        Si = S[idx]
        K_minus = (
            Si[0] / max(S_ai[2], 1e-9),
            Si[1] / max(S_ai[1], 1e-9),
            Si[2] / max(S_ai[0], 1e-9),
        )
        K_plus = (
            Si[0] / max(S_id[2], 1e-9),
            Si[1] / max(S_id[1], 1e-9),
            Si[2] / max(S_id[0], 1e-9),
        )
        T_i = tfn_add(K_minus, K_plus)
        results.append({
            'Alternative': alt,
            'S_i': Si,
            'K_minus': K_minus,
            'K_plus': K_plus,
            'T_i': T_i,
        })

    df_k = pd.DataFrame({
        'Alternative': [r['Alternative'] for r in results],
        'S_i': [format_tfn(r['S_i']) for r in results],
        'Fuzzy K-': [format_tfn(r['K_minus']) for r in results],
        'Fuzzy K+': [format_tfn(r['K_plus']) for r in results],
        'T_i': [format_tfn(r['T_i']) for r in results],
        'Crisp T_i': [defuzz_marcos_tfn(r['T_i']) for r in results],
    })
    st.subheader('Step 7: Fuzzy Utility Degrees and Total Utility Degree')
    st.dataframe(df_k, use_container_width=True, hide_index=True)

    # =====================================================
    # Step 8: dfcrisp MAX
    # dfcrisp = MAX(crisp T_i), not max fuzzy component
    # =====================================================
    crisp_T = [defuzz_marcos_tfn(r['T_i']) for r in results]
    dfcrisp = max(crisp_T) if crisp_T else 1e-9
    st.metric('dfcrisp MAX', f'{dfcrisp:.6f}')

    # =====================================================
    # Step 9-10: Fuzzy F(K-), F(K+), crisp utilities and ranking
    # F(K-) = K+ / dfcrisp
    # F(K+) = K- / dfcrisp
    # =====================================================
    final_rows = []
    for r in results:
        fuzzy_f_k_minus = tfn_scalar_div(r['K_plus'], dfcrisp)
        fuzzy_f_k_plus = tfn_scalar_div(r['K_minus'], dfcrisp)

        crisp_k_minus = defuzz_marcos_tfn(r['K_minus'])
        crisp_k_plus = defuzz_marcos_tfn(r['K_plus'])
        crisp_f_k_minus = defuzz_marcos_tfn(fuzzy_f_k_minus)
        crisp_f_k_plus = defuzz_marcos_tfn(fuzzy_f_k_plus)

        term_minus = ((1 - crisp_f_k_minus) / crisp_f_k_minus) if crisp_f_k_minus > 0 else 0.0
        term_plus = ((1 - crisp_f_k_plus) / crisp_f_k_plus) if crisp_f_k_plus > 0 else 0.0
        denom = 1 + term_minus + term_plus
        utility = (crisp_k_minus + crisp_k_plus) / denom if denom != 0 else 0.0

        final_rows.append({
            'Alternative': r['Alternative'],
            'Fuzzy K-': format_tfn(r['K_minus']),
            'Fuzzy K+': format_tfn(r['K_plus']),
            'T_i': format_tfn(r['T_i']),
            'Fuzzy F(K-)': format_tfn(fuzzy_f_k_minus),
            'Fuzzy F(K+)': format_tfn(fuzzy_f_k_plus),
            'Crisp K-': crisp_k_minus,
            'Crisp K+': crisp_k_plus,
            'Crisp F(K-)': crisp_f_k_minus,
            'Crisp F(K+)': crisp_f_k_plus,
            '(1-f(K-))/f(K-)': term_minus,
            '(1-f(K+))/f(K+)': term_plus,
            'f(K)': utility,
        })

    final_df = pd.DataFrame(final_rows).sort_values('f(K)', ascending=False).reset_index(drop=True)
    final_df['Rank'] = range(1, len(final_df) + 1)

    st.subheader('Step 9–10: Final Utility Function and Ranking')
    st.dataframe(final_df, use_container_width=True, hide_index=True)

    best_alt = final_df.iloc[0]['Alternative']
    best_score = final_df.iloc[0]['f(K)']
    st.success(f'Best Alternative: {best_alt} with utility score {best_score:.6f}')

    export_df = final_df.rename(columns={
        'Crisp K-': 'Crisp Ki-',
        'Crisp K+': 'Crisp Ki+',
        'f(K)': 'f(Ki)'
    }).copy()

    all_data = {
        'n_alternatives': len(alternatives),
        'n_criteria': len(criteria),
        'n_experts': st.session_state.marcos_data['n_experts'],
        'final_results': export_df,
        'best_alternative': best_alt,
        'best_score': best_score,
    }

    doc_bytes = create_marcos_word_document(all_data)
    st.download_button(
        label='Export MARCOS Results to Word',
        data=doc_bytes,
        file_name='Triangular_Fuzzy_MARCOS_Results.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        use_container_width=True,
    )

    if st.button('Start Over', key='marcos_restart'):
        st.session_state.marcos_step = 1
        st.session_state.marcos_data = {}
        st.rerun()


def marcos_step7_calculations_old_unused():
    st.header('Step 7: Fuzzy MARCOS Results')

    alternatives = st.session_state.marcos_data['alternatives']
    criteria = st.session_state.marcos_data['criteria']
    fuzzy_matrix = st.session_state.marcos_data['fuzzy_matrix']
    criteria_info = st.session_state.marcos_data['criteria_info']
    weights = criteria_info['Weight'].astype(float).values
    crit_types = criteria_info['Type'].tolist()

    n_alt = len(alternatives)
    n_crit = len(criteria)

    # Step 1-2: Expanded matrix with AAI and AI
    aai = []
    ai = []
    for j in range(n_crit):
        col = [fuzzy_matrix[i][j] for i in range(n_alt)]
        if crit_types[j] == 'Benefit':
            aai.append((min(x[0] for x in col), min(x[1] for x in col), min(x[2] for x in col)))
            ai.append((max(x[0] for x in col), max(x[1] for x in col), max(x[2] for x in col)))
        else:
            aai.append((max(x[0] for x in col), max(x[1] for x in col), max(x[2] for x in col)))
            ai.append((min(x[0] for x in col), min(x[1] for x in col), min(x[2] for x in col)))

    st.subheader('Step 1–2: Anti-Ideal (AAI) and Ideal (AI) Solutions')
    df_ref = pd.DataFrame({
        'Criterion': criteria,
        'Type': crit_types,
        'AAI': [format_tfn(x) for x in aai],
        'AI': [format_tfn(x) for x in ai],
    })
    st.dataframe(df_ref, use_container_width=True, hide_index=True)

    # Step 3: Fuzzy normalized matrix
    norm = [[None for _ in range(n_crit)] for _ in range(n_alt + 2)]
    labels = ['AAI'] + alternatives + ['AI']
    expanded = [aai] + fuzzy_matrix + [ai]

    for i in range(n_alt + 2):
        for j in range(n_crit):
            x = expanded[i][j]
            if crit_types[j] == 'Benefit':
                denom = ai[j][2]
                denom = denom if abs(denom) > 1e-9 else 1e-9
                norm[i][j] = (x[0] / denom, x[1] / denom, x[2] / denom)
            else:
                a = ai[j][0]
                norm[i][j] = (
                    a / max(x[2], 1e-9),
                    a / max(x[1], 1e-9),
                    a / max(x[0], 1e-9),
                )

    st.subheader('Step 3: Fuzzy Normalized Matrix')
    norm_df = pd.DataFrame(index=labels, columns=criteria)
    for i, lab in enumerate(labels):
        for j, c in enumerate(criteria):
            norm_df.loc[lab, c] = format_tfn(norm[i][j])
    st.dataframe(norm_df, use_container_width=True)

    # Step 4: Weighted normalized matrix
    weighted = [[None for _ in range(n_crit)] for _ in range(n_alt + 2)]
    for i in range(n_alt + 2):
        for j in range(n_crit):
            weighted[i][j] = (
                norm[i][j][0] * weights[j],
                norm[i][j][1] * weights[j],
                norm[i][j][2] * weights[j],
            )

    st.subheader('Step 4: Weighted Fuzzy Normalized Matrix')
    weighted_df = pd.DataFrame(index=labels, columns=criteria)
    for i, lab in enumerate(labels):
        for j, c in enumerate(criteria):
            weighted_df.loc[lab, c] = format_tfn(weighted[i][j])
    st.dataframe(weighted_df, use_container_width=True)

    # Step 5: S_i
    S = []
    for i in range(n_alt + 2):
        s = (0.0, 0.0, 0.0)
        for j in range(n_crit):
            s = tfn_add(s, weighted[i][j])
        S.append(s)

    s_df = pd.DataFrame({
        'Alternative': labels,
        'S_i (TFN)': [format_tfn(x) for x in S],
    })
    st.subheader('Step 5: Total Weighted Values (S_i)')
    st.dataframe(s_df, use_container_width=True, hide_index=True)

    S_aai = S[0]
    S_ai = S[-1]

    # Step 6: K- and K+
    results = []
    for idx, alt in enumerate(alternatives, start=1):
        Si = S[idx]
        K_minus = (Si[0] / max(S_aai[2], 1e-9), Si[1] / max(S_aai[1], 1e-9), Si[2] / max(S_aai[0], 1e-9))
        K_plus = (Si[0] / max(S_ai[2], 1e-9), Si[1] / max(S_ai[1], 1e-9), Si[2] / max(S_ai[0], 1e-9))
        t_i = tfn_add(K_minus, K_plus)
        results.append({
            'Alternative': alt,
            'S_i': Si,
            'K_minus': K_minus,
            'K_plus': K_plus,
            't_i': t_i,
        })

    df_k = pd.DataFrame({
        'Alternative': [r['Alternative'] for r in results],
        'Fuzzy K-': [format_tfn(r['K_minus']) for r in results],
        'Fuzzy K+': [format_tfn(r['K_plus']) for r in results],
        't_i': [format_tfn(r['t_i']) for r in results],
    })
    st.subheader('Step 6–7: Utility Degrees and Total Utility Degree')
    st.dataframe(df_k, use_container_width=True, hide_index=True)

    # Step 7-8: d and dfcrisp
    d = (
        max(r['t_i'][0] for r in results),
        max(r['t_i'][1] for r in results),
        max(r['t_i'][2] for r in results),
    )
    dfcrisp = defuzz_tfn(d)
    st.metric('dfcrisp', f'{dfcrisp:.6f}')

    # Step 8-9: f(K-), f(K+), final utility
    final_rows = []
    for r in results:
        fuzzy_f_k_minus = tfn_scalar_div(r['K_plus'], dfcrisp)
        fuzzy_f_k_plus = tfn_scalar_div(r['K_minus'], dfcrisp)

        crisp_k_minus = defuzz_tfn(r['K_minus'])
        crisp_k_plus = defuzz_tfn(r['K_plus'])
        crisp_f_k_minus = defuzz_tfn(fuzzy_f_k_minus)
        crisp_f_k_plus = defuzz_tfn(fuzzy_f_k_plus)

        term_minus = ((1 - crisp_f_k_minus) / crisp_f_k_minus) if crisp_f_k_minus > 0 else 0.0
        term_plus = ((1 - crisp_f_k_plus) / crisp_f_k_plus) if crisp_f_k_plus > 0 else 0.0
        denom = 1 + term_minus + term_plus
        utility = (crisp_k_plus + crisp_k_minus) / denom if denom != 0 else 0.0

        final_rows.append({
            'Alternative': r['Alternative'],
            'Fuzzy K-': format_tfn(r['K_minus']),
            'Fuzzy K+': format_tfn(r['K_plus']),
            'Fuzzy F(K-)': format_tfn(fuzzy_f_k_minus),
            'Fuzzy F(K+)': format_tfn(fuzzy_f_k_plus),
            'Crisp Ki-': crisp_k_minus,
            'Crisp Ki+': crisp_k_plus,
            'Crisp F(K-)': crisp_f_k_minus,
            'Crisp F(K+)': crisp_f_k_plus,
            'f(Ki)': utility,
        })

    final_df = pd.DataFrame(final_rows).sort_values('f(Ki)', ascending=False).reset_index(drop=True)
    final_df['Rank'] = range(1, len(final_df) + 1)

    st.subheader('Step 8–10: Final Utility Function and Ranking')
    st.dataframe(final_df, use_container_width=True, hide_index=True)

    best_alt = final_df.iloc[0]['Alternative']
    best_score = final_df.iloc[0]['f(Ki)']
    st.success(f'Best Alternative: {best_alt} with utility score {best_score:.6f}')

    all_data = {
        'n_alternatives': len(alternatives),
        'n_criteria': len(criteria),
        'n_experts': st.session_state.marcos_data['n_experts'],
        'final_results': final_df,
        'best_alternative': best_alt,
        'best_score': best_score,
    }

    doc_bytes = create_marcos_word_document(all_data)
    st.download_button(
        label='Export MARCOS Results to Word',
        data=doc_bytes,
        file_name='Triangular_Fuzzy_MARCOS_Results.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        use_container_width=True,
    )

    if st.button('Start Over', key='marcos_restart'):
        st.session_state.marcos_step = 1
        st.session_state.marcos_data = {}
        st.rerun()


# =========================================================
# MAIN
# =========================================================
def main():
    st.sidebar.title('Fuzzy MCDM Model Selection')
    st.sidebar.markdown('Select the model you want to use:')

    model_choice = st.sidebar.radio(
        'Choose Model:',
        ['Triangular Fuzzy OPA', 'Triangular Fuzzy MARCOS Method'],
        index=0,
    )

    if model_choice == 'Triangular Fuzzy OPA':
        opa_model()
    else:
        marcos_model()

    st.markdown(
        """
        <div class="footer">
        <p>Integrated Fuzzy MCDM Models | OPA + Fuzzy MARCOS for Moktadir, M. A.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


if __name__ == '__main__':
    main()
